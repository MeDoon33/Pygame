"""
Dice Roguelite - Pygame Game
Tựa game giống Dicero: nhân vật chạy tự động, hệ thống xúc xắc poker
"""

import pygame
import random
import math
import sys
import json
import os
from enum import Enum
from tkinter import filedialog, Tk

# --- TILEMAP SYSTEM ---
from tilemap import MapManager

# --- ANIMATION SYSTEM ---
from enemy_attack_animation import EnemyAttackAnimator
from enemy_projectile_animation import Projectile

pygame.init()
pygame.mixer.init()


def load_sprite(path, width, height):
    try:
        img = pygame.image.load(path).convert_alpha()
        return pygame.transform.scale(img, (width, height))
    except:
        # Nếu không tìm thấy file, tạo một bề mặt tạm thời màu tím để báo lỗi
        surf = pygame.Surface((width, height), pygame.SRCALPHA)
        pygame.draw.rect(surf, (255, 0, 255), (0, 0, width, height), 2)
        return surf
# === CẤU HÌNH ===
SCREEN_W, SCREEN_H = 480, 820
FPS = 60

# Màu sắc
C_BG        = (13, 10, 26)
C_PATH      = (42, 16, 96)
C_PATH2     = (55, 25, 120)
C_WHITE     = (255, 255, 255)
C_GOLD      = (255, 215, 0)
C_RED       = (220, 50, 50)
C_CRIMSON   = (180, 20, 40)
C_PURPLE    = (120, 60, 200)
C_LAVENDER  = (170, 140, 255)
C_DARK      = (20, 10, 40)
C_GRAY      = (80, 80, 100)
C_HP_RED    = (220, 60, 60)
C_HP_BG     = (40, 20, 60)
C_ORANGE    = (255, 140, 30)
C_CYAN      = (80, 220, 255)
C_GREEN     = (60, 200, 80)

screen = pygame.display.set_mode((SCREEN_W, SCREEN_H))
pygame.display.set_caption("⚔️  Dice Roguelite")
clock = pygame.time.Clock()

# Font
try:
    font_lg   = pygame.font.SysFont("segoeui", 28, bold=True)
    font_md   = pygame.font.SysFont("segoeui", 20, bold=True)
    font_sm   = pygame.font.SysFont("segoeui", 15, bold=True)
    font_xs   = pygame.font.SysFont("segoeui", 13)
    font_die  = pygame.font.SysFont("segoeui", 30, bold=True)
    font_hand = pygame.font.SysFont("segoeui", 17, bold=True)
except:
    font_lg = font_md = font_sm = font_xs = font_die = font_hand = pygame.font.Font(None, 24)


# ==================== UTILITIES ====================

def lerp(a, b, t):
    return a + (b - a) * t

def draw_rounded_rect(surf, color, rect, radius, border=0, border_color=None):
    pygame.draw.rect(surf, color, rect, border_radius=radius)
    if border and border_color:
        pygame.draw.rect(surf, border_color, rect, border, border_radius=radius)

def draw_text_center(surf, text, font, color, cx, cy, shadow=True):
    if shadow:
        s = font.render(text, True, (0, 0, 0))
        surf.blit(s, s.get_rect(center=(cx+2, cy+2)))
    t = font.render(text, True, color)
    surf.blit(t, t.get_rect(center=(cx, cy)))

def draw_text(surf, text, font, color, x, y, shadow=True):
    if shadow:
        s = font.render(text, True, (0, 0, 0))
        surf.blit(s, (x+1, y+1))
    t = font.render(text, True, color)
    surf.blit(t, (x, y))

def load_spritesheet(path, frame_count, width, height):
    """Cắt một dải ảnh ngang thành danh sách các frames."""
    frames = []
    try:
        sheet = pygame.image.load(path).convert_alpha()
        sheet_width = sheet.get_width()
        frame_width = sheet_width // frame_count  # Tính width thực tế cho mỗi frame
        for i in range(frame_count):
            # Cắt từng khung hình dựa trên width tính toán
            frame = sheet.subsurface(pygame.Rect(i * frame_width, 0, frame_width, height))
            frames.append(pygame.transform.scale(frame, (width, height)))
    except Exception as e:
        print(f"Lỗi load spritesheet {path}: {e}")
        # Nếu lỗi, tạo khung hình mặc định trong suốt để tránh crash game
        for _ in range(frame_count):
            s = pygame.Surface((width, height), pygame.SRCALPHA)
            pygame.draw.line(s, (255, 0, 255), (4, 4), (width - 5, height - 5), 3)
            pygame.draw.line(s, (255, 0, 255), (width - 5, 4), (4, height - 5), 3)
            frames.append(s)
    return frames


def load_spritesheet_vertical(path, frame_count, width, height):
    """Cắt một dải ảnh dọc thành danh sách các frames (frame xếp từ trên xuống dưới)."""
    frames = []
    try:
        sheet = pygame.image.load(path).convert_alpha()
        sheet_height = sheet.get_height()
        frame_height = sheet_height // frame_count  # Tính height thực tế cho mỗi frame
        for i in range(frame_count):
            # Cắt từng khung hình dựa trên height tính toán (từ trên xuống dưới)
            frame = sheet.subsurface(pygame.Rect(0, i * frame_height, width, frame_height))
            frames.append(pygame.transform.scale(frame, (width, height)))
    except Exception as e:
        print(f"Lỗi load spritesheet dọc {path}: {e}")
        # Nếu lỗi, tạo khung hình mặc định trong suốt để tránh crash game
        for _ in range(frame_count):
            s = pygame.Surface((width, height), pygame.SRCALPHA)
            pygame.draw.line(s, (255, 0, 255), (4, 4), (width - 5, height - 5), 3)
            pygame.draw.line(s, (255, 0, 255), (width - 5, 4), (4, height - 5), 3)
            frames.append(s)
    return frames


def load_questions():
    """Load câu hỏi từ file JSON."""
    try:
        with open("questions.json", "r", encoding="utf-8") as f:
            data = json.load(f)
            return data.get("questions", [])
    except Exception as e:
        print(f"Lỗi load questions: {e}")
        return []


def save_questions(questions):
    """Lưu câu hỏi vào file JSON."""
    try:
        with open("questions.json", "w", encoding="utf-8") as f:
            json.dump({"questions": questions}, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"Lỗi save questions: {e}")
        return False


def open_file_chooser():
    """Mở file dialog để chọn file câu hỏi từ máy.
    
    Hỗ trợ các định dạng:
    - Text files (.txt)
    - Word documents (.doc, .docx)
    - Tất cả các file
    """
    try:
        root = Tk()
        root.withdraw()  # Ẩn window chính của tkinter
        file_path = filedialog.askopenfilename(
            title="Chọn file câu hỏi (.txt, .doc, .docx)",
            filetypes=[
                ("Text files", "*.txt"),
                ("Word documents", "*.doc"),
                ("Word documents (2007+)", "*.docx"),
                ("All files", "*.*")
            ],
            initialdir="."
        )
        root.destroy()
        return file_path
    except Exception as e:
        print(f"Lỗi mở file dialog: {e}")
        return None


def is_file_empty(file_path: str) -> bool:
    """Kiểm tra xem file có rỗng không.
    
    Args:
        file_path: Đường dẫn file
        
    Returns:
        True nếu file rỗng, False nếu có nội dung
    """
    try:
        file_size = os.path.getsize(file_path)
        return file_size == 0
    except Exception:
        return True


def is_text_file(file_path: str) -> bool:
    """Kiểm tra xem file có phải là file văn bản không.
    
    Hỗ trợ:
    - File .txt (UTF-8, ASCII, Latin-1, cp1252, iso-8859-1)
    - File .doc (Word 97-2003)
    - File .docx (Word 2007+)
    
    Args:
        file_path: Đường dẫn file
        
    Returns:
        True nếu là file văn bản, False nếu là file nhị phân khác
    """
    try:
        # Lấy phần mở rộng file
        _, file_ext = os.path.splitext(file_path)
        file_ext = file_ext.lower()
        
        # Kiểm tra loại file từ phần mở rộng
        if file_ext in ['.txt']:
            # Cố gắng đọc một số byte đầu để kiểm tra encoding
            with open(file_path, 'rb') as f:
                sample = f.read(1024)
                
            # Thử các encoding phổ biến
            encodings_to_try = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings_to_try:
                try:
                    sample.decode(encoding)
                    return True
                except UnicodeDecodeError:
                    continue
            
            # Nếu không decode được với strict mode, có thể vẫn là text
            # Thử với errors='ignore'
            try:
                decoded = sample.decode('utf-8', errors='ignore')
                if decoded.strip():  # Có nội dung text
                    return True
            except:
                pass
                
            return False
                        
        elif file_ext in ['.doc', '.docx']:
            # Word documents được coi là file văn bản hợp lệ
            return True
        else:
            # Cố gắng đọc file để kiểm tra
            with open(file_path, 'rb') as f:
                sample = f.read(512)
                
            # Kiểm tra xem có là text không
            encodings_to_try = ['utf-8', 'latin-1', 'cp1252']
            for encoding in encodings_to_try:
                try:
                    sample.decode(encoding)
                    return True
                except UnicodeDecodeError:
                    continue
            
            return False
    except Exception:
        return False


def extract_text_from_doc(file_path: str) -> str:
    """Trích xuất text từ file .doc (Word 97-2003).
    
    Args:
        file_path: Đường dẫn file .doc
        
    Returns:
        Nội dung text từ file hoặc rỗng nếu lỗi
    """
    try:
        # Thử sử dụng python-docx nếu có cài đặt (chỉ cho .docx)
        import subprocess
        result = subprocess.run(
            [sys.executable, '-m', 'pip', 'show', 'python-docx'],
            capture_output=True,
            text=True
        )
        if result.returncode == 0:
            # python-docx được cài đặt, sử dụng nó cho .docx
            try:
                from docx import Document
                doc = Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                return text
            except:
                pass
        
        # Fallback: Sử dụng file reading thử - chỉ cho .txt
        # Đây là fallback nếu không có thư viện
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        print(f"Lỗi trích xuất text từ .doc: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Trích xuất text từ file .docx (Word 2007+).
    
    Args:
        file_path: Đường dẫn file .docx
        
    Returns:
        Nội dung text từ file hoặc rỗng nếu lỗi
    """
    try:
        from docx import Document
        doc = Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except ImportError:
        print("[WARN] python-docx library is not installed.")
        print("   Install it with: pip install python-docx")
        return ""
    except Exception as e:
        print(f"Lỗi trích xuất text từ .docx: {e}")
        return ""


def read_text_content(file_path: str) -> tuple[bool, str]:
    """Đọc nội dung text từ file.
    
    Hỗ trợ: .txt, .doc, .docx
    Thử nhiều encoding khác nhau để đọc file .txt
    
    Args:
        file_path: Đường dẫn file
        
    Returns:
        Tuple (success: bool, content: str)
        - success: True nếu đọc thành công
        - content: Nội dung text hoặc thông báo lỗi
    """
    try:
        _, file_ext = os.path.splitext(file_path)
        file_ext = file_ext.lower()
        
        if file_ext == '.txt':
            # Thử nhiều encoding cho file .txt
            encodings_to_try = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='strict') as f:
                        content = f.read()
                    # Kiểm tra xem có phải là encoding đúng không (không có ký tự lạ)
                    if '\ufffd' not in content:  # \ufffd là replacement character
                        return True, content
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            # Nếu tất cả encoding đều fail, thử với errors='ignore'
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                if content.strip():  # Có nội dung
                    return True, content
            except Exception:
                pass
            
            return False, "Không thể đọc file .txt với bất kỳ encoding nào được hỗ trợ"
            
        elif file_ext == '.docx':
            content = extract_text_from_docx(file_path)
            if content:
                return True, content
            else:
                return False, "Không thể trích xuất text từ file .docx"
                
        elif file_ext == '.doc':
            content = extract_text_from_doc(file_path)
            if content:
                return True, content
            else:
                return False, "Không thể trích xuất text từ file .doc"
        else:
            # Cố gắng đọc file khác
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                return True, content
            except Exception as e:
                return False, f"Lỗi đọc file: {str(e)}"
            
    except Exception as e:
        return False, f"Lỗi đọc file: {str(e)}"


def import_questions_from_file(file_path):
    """
    Import câu hỏi từ file văn bản hoặc Word document.
    
    Hỗ trợ:
    - File .txt (UTF-8, ASCII, Latin-1)
    - File .docx (Word 2007+)
    - File .doc (Word 97-2003, cần python-docx)
    
    Format:
    Câu hỏi?
    Option 1
    Option 2
    Option 3
    Option 4
    Correct: 1
    Reward: Tên reward   # Tùy chọn, sẽ bị bỏ qua khi import
    
    [blank line] hoặc không có blank line giữa các câu hỏi
    
    Câu hỏi tiếp theo?
    ...
    
    Returns:
        Danh sách câu hỏi (list[dict]) hoặc rỗng nếu lỗi
    """
    questions = []
    
    # ===== KIỂM TRA 1: File có rỗng không? =====
    if is_file_empty(file_path):
        print("[ERROR] File is empty. Please select a non-empty file.")
        return []
    
    # ===== KIỂM TRA 2: File có phải văn bản không? =====
    if not is_text_file(file_path):
        print("[ERROR] File is not a supported text file. Supports .txt, .doc, .docx.")
        return []
    
    # ===== ĐỌC NỘI DUNG FILE =====
    success, content = read_text_content(file_path)
    if not success:
        print(f"[ERROR] {content}")
        return []
    
    if not content or not content.strip():
        print("[ERROR] File does not contain valid text content.")
        return []
    
    # ===== XỬ LÝ NỘI DUNG =====
    try:
        content = content.replace("\r\n", "\n").replace("\r", "\n")
        lines = [line.strip() for line in content.split("\n")]
        next_id = 1
        idx = 0
        
        while idx < len(lines):
            # Skip empty lines between question blocks
            while idx < len(lines) and not lines[idx]:
                idx += 1
            if idx >= len(lines):
                break
            
            question_text = lines[idx]
            idx += 1
            if not question_text:
                continue
            
            # Extract exactly 4 answer options
            options = []
            while idx < len(lines) and len(options) < 4:
                if lines[idx].lower().startswith("correct:"):
                    break
                if lines[idx]:
                    options.append(lines[idx])
                idx += 1
            
            if len(options) < 4:
                continue
            
            # Extract correct answer
            correct_idx = 0
            while idx < len(lines):
                if lines[idx].lower().startswith("correct:"):
                    try:
                        correct_idx = int(lines[idx].split(":", 1)[1].strip()) - 1
                        if correct_idx < 0 or correct_idx >= 4:
                            correct_idx = 0
                    except Exception:
                        correct_idx = 0
                    idx += 1
                    break
                idx += 1

            # Skip optional reward lines if present
            while idx < len(lines) and lines[idx].lower().startswith("reward:"):
                idx += 1
            
            questions.append({
                "id": next_id,
                "text": question_text,
                "options": options,
                "correct": correct_idx
            })
            next_id += 1
        
        if not questions:
            print("[ERROR] No valid questions found in file.")
            print("   Check file format against the sample file.")
            return []
        
        print(f"[OK] Imported {len(questions)} questions.")
        return questions
    except Exception as e:
        print(f"[ERROR] Import questions error: {e}")
        return []


# ==================== PARTICLES ====================

class Particle:
    def __init__(self, x, y, color, size=5, vx=None, vy=None, life=None):
        self.x = x
        self.y = y
        self.color = color
        self.size = size
        self.vx = vx if vx is not None else random.uniform(-3, 3)
        self.vy = vy if vy is not None else random.uniform(-4, -1)
        self.life = life or random.randint(30, 50)
        self.max_life = self.life

    def update(self):
        self.x += self.vx
        self.y += self.vy
        self.vy += 0.15
        self.size *= 0.97
        self.life -= 1

    def draw(self, surf):
        alpha = int(255 * self.life / self.max_life)
        s = pygame.Surface((int(self.size*2+2), int(self.size*2+2)), pygame.SRCALPHA)
        pygame.draw.circle(s, (*self.color, alpha), (int(self.size)+1, int(self.size)+1), max(1, int(self.size)))
        surf.blit(s, (int(self.x - self.size), int(self.y - self.size)))


class FloatingText:
    def __init__(self, x, y, text, color, size=22):
        self.x = x
        self.y = y
        self.text = text
        self.color = color
        self.size = size
        self.vy = -2.0
        self.life = 70
        self.max_life = 70
        try:
            self.font = pygame.font.SysFont("segoeui", size, bold=True)
        except:
            self.font = pygame.font.Font(None, size)

    def update(self):
        self.y += self.vy
        self.vy *= 0.96
        self.life -= 1

    def draw(self, surf):
        alpha = int(255 * self.life / self.max_life)
        t = self.font.render(self.text, True, self.color)
        s = pygame.Surface(t.get_size(), pygame.SRCALPHA)
        s.fill((0, 0, 0, 0))
        s.blit(t, (0, 0))
        s.set_alpha(alpha)
        surf.blit(s, (int(self.x - t.get_width()//2), int(self.y)))


# ==================== DICE ====================

FACE_DOTS = {
    1: [(0, 0)],
    2: [(-1, -1), (1, 1)],
    3: [(-1, -1), (0, 0), (1, 1)],
    4: [(-1, -1), (1, -1), (-1, 1), (1, 1)],
    5: [(-1, -1), (1, -1), (0, 0), (-1, 1), (1, 1)],
    6: [(-1, -1), (0, -1), (1, -1), (-1, 1), (0, 1), (1, 1)],
}

# ── pre-render each face as a Surface so we can scale/transform quickly ──
_FACE_CACHE = {}

def _build_face(value, size, bg, border, dot_color=(255,255,255)):
    """Render a single die face onto a Surface of given size."""
    s = pygame.Surface((size, size), pygame.SRCALPHA)
    r = pygame.Rect(0, 0, size, size)
    radius = max(4, size // 6)
    pygame.draw.rect(s, bg, r, border_radius=radius)
    pygame.draw.rect(s, border, r, 2, border_radius=radius)
    # inner bevel shadow
    inner = r.inflate(-4, -4)
    pygame.draw.rect(s, (20, 10, 40), inner, border_radius=max(2, radius-2))
    pygame.draw.rect(s, bg, inner.inflate(-2, -2), border_radius=max(1, radius-3))

    cx, cy   = size // 2, size // 2
    offset   = size // 4
    dot_r    = max(3, size // 12)
    dots     = FACE_DOTS.get(value, [])
    for (dx, dy) in dots:
        px = cx + dx * offset
        py = cy + dy * offset
        # soft glow halo
        glow_r = dot_r * 2
        g = pygame.Surface((glow_r*2, glow_r*2), pygame.SRCALPHA)
        pygame.draw.circle(g, (*dot_color, 55), (glow_r, glow_r), glow_r)
        s.blit(g, (px - glow_r, py - glow_r))
        pygame.draw.circle(s, dot_color, (px, py), dot_r)
    return s


def _get_face(value, size=60, locked=False):
    key = (value, size, locked)
    if key not in _FACE_CACHE:
        if locked:
            bg, border = (60, 45, 10), C_GOLD
        else:
            bg, border = (40, 20, 75), (100, 60, 160)
        _FACE_CACHE[key] = _build_face(value, size, bg, border)
    return _FACE_CACHE[key]


# ── Die trail ghost ──────────────────────────────────────────────────────
class DieTrail:
    """A fading ghost copy left behind during roll flight."""
    def __init__(self, x, y, angle, value, size):
        self.x, self.y  = x, y
        self.angle      = angle
        self.value      = value
        self.size       = size
        self.life       = 18
        self.max_life   = 18

    def update(self):
        self.life -= 1

    def draw(self, surf):
        if self.life <= 0:
            return
        alpha = int(160 * self.life / self.max_life)
        face  = _get_face(self.value, self.size)
        rot   = pygame.transform.rotate(face, self.angle)
        rot_a = rot.copy()
        rot_a.set_alpha(alpha)
        surf.blit(rot_a, rot_a.get_rect(center=(int(self.x), int(self.y))))


# ── Main Die class ───────────────────────────────────────────────────────
class Die:
    SIZE   = 60
    RADIUS = 10

    # Roll animation phases
    PHASE_IDLE    = 0   # sitting still, empty
    PHASE_FLYING  = 1   # launched from off-screen, arcing in
    PHASE_SPIN    = 2   # tumbling fast on the table after landing
    PHASE_BOUNCE  = 3   # small bounce settle
    PHASE_LAND    = 4   # flash + done
    PHASE_LOCKED  = 5   # gold locked

    TOTAL_ROLL_FRAMES = 40   # flight+spin total

    def __init__(self, index):
        self.index   = index
        self.value   = 1
        self.locked  = False
        self.empty   = True
        self.hover   = False

        # animation state
        self.phase       = Die.PHASE_IDLE
        self.anim_t      = 0      # counts up each frame while animating
        self.spin_face   = 1      # rapid face cycling during spin
        self.spin_angle  = 0.0   # visual rotation angle (degrees)
        self.spin_speed  = 0.0   # deg/frame
        self.scale       = 1.0   # for bounce squash
        self.flash_alpha = 0     # white flash on land
        self.land_glow   = 0     # purple glow frames after landing
        self.combo_glow  = False
        self.combo_color = None

        # flight arc (PHASE_FLYING)
        self.fly_start_x = 0.0
        self.fly_start_y = 0.0
        self.fly_peak_y  = 0.0   # apex of arc

        # target resting position (set by layout each frame)
        self.x = 0
        self.y = 0

        # ghost trails
        self.trails: list[DieTrail] = []
        self._trail_timer = 0

    # ── public ──────────────────────────────────────
    def roll(self, stagger_frames: int = 0):
        """Start a full cinematic roll animation."""
        self.value     = random.randint(1, 6)
        self.empty     = False
        self.locked    = False
        self.spin_face = random.randint(1, 6)
        self.spin_speed = random.uniform(18, 28)   # deg/frame
        self.spin_angle = random.uniform(0, 360)
        self.scale      = 1.0
        self.flash_alpha = 0
        self.land_glow   = 0
        self.combo_glow  = False
        self.combo_color = None
        self.trails.clear()
        self.anim_t      = -stagger_frames          # negative = waiting
        self.phase       = Die.PHASE_FLYING

        # launch from a random off-screen edge (top or left or right)
        side = random.choice(['top', 'left', 'right'])
        if side == 'top':
            self.fly_start_x = self.x + random.uniform(-80, 80)
            self.fly_start_y = self.y - random.uniform(200, 340)
        elif side == 'left':
            self.fly_start_x = self.x - random.uniform(180, 320)
            self.fly_start_y = self.y + random.uniform(-40, 40)
        else:
            self.fly_start_x = self.x + random.uniform(180, 320)
            self.fly_start_y = self.y + random.uniform(-40, 40)
        self.fly_peak_y = min(self.fly_start_y, self.y) - random.uniform(30, 70)

    def get_rect(self):
        return pygame.Rect(
            int(self.x - self.SIZE * self.scale / 2),
            int(self.y - self.SIZE * self.scale / 2),
            int(self.SIZE * self.scale),
            int(self.SIZE * self.scale),
        )

    # ── internal helpers ────────────────────────────
    def _fly_pos(self, t01):
        """Quadratic bezier arc: start -> peak-control -> target."""
        # control point is horizontally between start & end, at peak height
        cx = (self.fly_start_x + self.x) / 2
        cy = self.fly_peak_y
        bx = (1-t01)**2 * self.fly_start_x + 2*(1-t01)*t01*cx + t01**2 * self.x
        by = (1-t01)**2 * self.fly_start_y + 2*(1-t01)*t01*cy + t01**2 * self.y
        return bx, by

    def _ease_out_bounce(self, t):
        """Easing used during SPIN to slow to a stop."""
        if t < 1/2.75:
            return 7.5625 * t * t
        elif t < 2/2.75:
            t -= 1.5/2.75
            return 7.5625 * t * t + 0.75
        elif t < 2.5/2.75:
            t -= 2.25/2.75
            return 7.5625 * t * t + 0.9375
        else:
            t -= 2.625/2.75
            return 7.5625 * t * t + 0.984375

    # ── update ──────────────────────────────────────
    def update(self):
        # update ghost trails
        for tr in self.trails:
            tr.update()
        self.trails = [tr for tr in self.trails if tr.life > 0]

        if self.flash_alpha > 0:
            self.flash_alpha = max(0, self.flash_alpha - 18)
        if self.land_glow > 0:
            self.land_glow -= 1

        if self.phase == Die.PHASE_IDLE or self.phase == Die.PHASE_LOCKED:
            return

        self.anim_t += 1
        if self.anim_t <= 0:          # stagger delay
            return

        FLY_DUR  = 22   # frames for flight arc
        SPIN_DUR = 20   # frames for table spin after landing
        BOUNCE_DUR = 8  # settle bounce

        # ── FLYING ──────────────────────────────────
        if self.phase == Die.PHASE_FLYING:
            t01 = min(1.0, self.anim_t / FLY_DUR)
            bx, by = self._fly_pos(t01)

            # squash during arc: elongate while flying, squash on impact
            if t01 < 0.8:
                self.scale = 1.0 + 0.15 * math.sin(t01 * math.pi)
            else:
                self.scale = 1.0

            # rapid face cycling
            if self.anim_t % 3 == 0:
                self.spin_face = random.randint(1, 6)

            # rotation: fast spin
            self.spin_angle += self.spin_speed

            # spawn trail ghost every 3 frames
            self._trail_timer += 1
            if self._trail_timer >= 3:
                self._trail_timer = 0
                size = max(20, int(self.SIZE * self.scale * 0.85))
                self.trails.append(DieTrail(bx, by, self.spin_angle, self.spin_face, size))

            # store current draw position temporarily via x/y offset trick
            self._draw_x = bx
            self._draw_y = by

            if t01 >= 1.0:
                # landing impact
                self.phase   = Die.PHASE_SPIN
                self.anim_t  = 0
                self.scale   = 1.35   # squash wide on impact
                self._draw_x = self.x
                self._draw_y = self.y
                # spawn impact burst stored in game (we fire from draw)
                self._impact = True

        # ── SPIN (tumbling on table) ─────────────────
        elif self.phase == Die.PHASE_SPIN:
            t01 = min(1.0, self.anim_t / SPIN_DUR)
            # scale: squash -> overshoot -> settle
            if t01 < 0.15:
                self.scale = 1.35 - 0.35 * (t01 / 0.15)   # unsquash
            elif t01 < 0.6:
                self.scale = 1.0 + 0.08 * math.sin((t01 - 0.15) / 0.45 * math.pi)
            else:
                self.scale = 1.0

            # decelerate spin
            speed = self.spin_speed * (1.0 - t01) ** 1.5
            self.spin_angle += speed

            # face flicker: random early, slow down
            flip_prob = int(6 * (1.0 - t01)) + 1
            if self.anim_t % flip_prob == 0:
                self.spin_face = random.randint(1, 6)

            self._draw_x = self.x
            self._draw_y = self.y

            if t01 >= 1.0:
                self.phase      = Die.PHASE_BOUNCE
                self.anim_t     = 0
                self.spin_face  = self.value   # lock final face
                self.spin_angle = 0.0

        # ── BOUNCE (settle) ──────────────────────────
        elif self.phase == Die.PHASE_BOUNCE:
            t01 = min(1.0, self.anim_t / BOUNCE_DUR)
            # small vertical bounce
            bounce = math.sin(t01 * math.pi * 2) * 6 * (1.0 - t01)
            self._draw_y = self.y - bounce
            self._draw_x = self.x
            self.scale   = 1.0

            if t01 >= 1.0:
                self.phase       = Die.PHASE_LAND
                self.anim_t      = 0
                self.flash_alpha = 220
                self.land_glow   = 30
                self._draw_x     = self.x
                self._draw_y     = self.y

        # ── LAND ─────────────────────────────────────
        elif self.phase == Die.PHASE_LAND:
            self._draw_x = self.x
            self._draw_y = self.y
            if self.anim_t > 10:
                self.phase = Die.PHASE_IDLE

    def is_animating(self):
        return self.phase not in (Die.PHASE_IDLE, Die.PHASE_LOCKED)

    # ── draw ────────────────────────────────────────
    def draw(self, surf):
        # draw trails first (behind die)
        for tr in self.trails:
            tr.draw(surf)

        # ── EMPTY slot ──────────────────────────────
        if self.empty:
            r = self.get_rect()
            pygame.draw.rect(surf, (50, 25, 80), r, border_radius=self.RADIUS)
            pygame.draw.rect(surf, (100, 60, 140), r, 2, border_radius=self.RADIUS)
            return

        # draw position & scale
        if hasattr(self, '_draw_x'):
            draw_cx = self._draw_x
            draw_cy = self._draw_y
        else:
            draw_cx = self.x
            draw_cy = self.y

        draw_size = max(10, int(self.SIZE * self.scale))

        # choose which face value to display
        display_val = self.spin_face if self.phase in (Die.PHASE_FLYING, Die.PHASE_SPIN) else self.value

        # ── land glow halo ───────────────────────────
        if self.land_glow > 0:
            g_alpha = int(180 * self.land_glow / 30)
            g_r     = draw_size // 2 + 14
            glow_s  = pygame.Surface((g_r*2, g_r*2), pygame.SRCALPHA)
            pygame.draw.circle(glow_s, (160, 100, 255, g_alpha), (g_r, g_r), g_r)
            surf.blit(glow_s, (int(draw_cx) - g_r, int(draw_cy) - g_r))

        # ── locked gold glow ─────────────────────────
        if self.locked:
            g_r = draw_size // 2 + 8
            glow_s = pygame.Surface((g_r*2, g_r*2), pygame.SRCALPHA)
            pulse  = int(60 + 40 * math.sin(pygame.time.get_ticks() * 0.006))
            pygame.draw.circle(glow_s, (255, 215, 0, pulse), (g_r, g_r), g_r)
            surf.blit(glow_s, (int(draw_cx) - g_r, int(draw_cy) - g_r))

        # ── combo highlight glow ─────────────────────
        if self.combo_glow and self.combo_color is not None:
            glow_r = draw_size // 2 + 12
            glow_s = pygame.Surface((glow_r*2, glow_r*2), pygame.SRCALPHA)
            pygame.draw.circle(glow_s, (*self.combo_color, 90), (glow_r, glow_r), glow_r)
            pygame.draw.circle(glow_s, (*self.combo_color, 180), (glow_r, glow_r), glow_r, 2)
            surf.blit(glow_s, (int(draw_cx) - glow_r, int(draw_cy) - glow_r))

        # ── render face ──────────────────────────────
        face = _get_face(display_val, draw_size, locked=self.locked)

        # hover tint
        if self.hover and not self.locked and self.phase == Die.PHASE_IDLE:
            face = face.copy()
            tint = pygame.Surface(face.get_size(), pygame.SRCALPHA)
            tint.fill((80, 50, 160, 60))
            face.blit(tint, (0, 0))

        if self.spin_angle != 0.0:
            face = pygame.transform.rotate(face, self.spin_angle)

        face_rect = face.get_rect(center=(int(draw_cx), int(draw_cy)))
        surf.blit(face, face_rect)

        # ── white flash overlay on landing ───────────
        if self.flash_alpha > 0:
            flash = face.copy()
            flash.fill((255, 255, 255, self.flash_alpha), special_flags=pygame.BLEND_RGBA_MULT)
            surf.blit(flash, face_rect)

        # ── lock badge ───────────────────────────────
        if self.locked:
            badge_font = font_xs
            badge = badge_font.render("LOCK", True, C_GOLD)
            bx = int(draw_cx) - badge.get_width() // 2
            by = int(draw_cy) + draw_size // 2 - 4
            surf.blit(badge, (bx, by))


# ==================== POKER HANDS ====================

def count_values(dice_vals):
    c = {}
    for v in dice_vals:
        c[v] = c.get(v, 0) + 1
    return c

def eval_hand(dice_vals):
    if not dice_vals:
        return "—", 0, 1.0, 0, None
    vals = sorted(dice_vals)
    c = count_values(vals)
    counts = sorted(c.values(), reverse=True)
    unique = len(c)
    base = sum(vals)

    combo_count = max((count for count in counts if count >= 3), default=0)
    combo_value = None
    if combo_count >= 3:
        combo_value = max(v for v, count in c.items() if count == combo_count)

    is_straight = (len(vals) >= 5 and unique == 5 and vals[-1] - vals[0] == 4)

    if counts[0] >= 5:
        return "Five of a Kind", base, 8.0, combo_count, combo_value
    if is_straight and counts[0] == 1 and len(vals) == 5:
        return "Royal Flush", base, 7.0, combo_count, combo_value
    if counts[0] >= 4:
        return "Four of a Kind", base, 5.0, combo_count, combo_value
    if unique == 2 and counts[0] == 3:
        return "Full House", base, 4.0, combo_count, combo_value
    if is_straight:
        return "Straight", base, 3.0, combo_count, combo_value
    if counts[0] >= 3:
        return "Three of a Kind", base, 2.5, combo_count, combo_value
    if unique <= len(vals) - 2 and counts[0] >= 2 and counts[1] >= 2:
        return "Two Pair", base, 2.0, combo_count, combo_value
    if counts[0] >= 2:
        return "One Pair", base, 1.5, combo_count, combo_value
    return "High Card", base, 1.0, combo_count, combo_value

HAND_COLORS = {
    "Five of a Kind": (255, 100, 255),
    "Royal Flush":    (255, 200, 50),
    "Four of a Kind": (255, 140, 50),
    "Full House":     (100, 220, 255),
    "Straight":       (80, 255, 180),
    "Three of a Kind":(180, 130, 255),
    "Two Pair":       (200, 200, 100),
    "One Pair":       (180, 180, 180),
    "High Card":      (120, 120, 120),
    "—":              (80, 80, 80),
}


# ==================== ENEMY ====================

class Enemy:
    def __init__(self, level=1, boss=False):
        self.boss = boss
        self.max_hp = 100 + level * 50 if boss else 60 + level * 25
        self.hp = self.max_hp
        self.x = SCREEN_W // 2
        self.y = 200
        self.frame = 0  # For bob/idle animation
        self.hit_timer = 0
        self.level = level
        self.atk = 6 + level * 6 if boss else 4+ level * 3
        self.dead = False
        self.death_timer = 0
        
        # Debuffs for boss
        self.stunned = False
        self.damage_reduction = 1.0
        self.burn_damage = 0
        self.burn_tick = 0
        self.bleed_mult = 1.0
        self.debuff_timers = {}
        
        # --- NẠP ANIMATION FRAMES ---
        self.sprite_width = 146
        self.sprite_height = 150
        self.frame_timer = 0
        self.current_frame = 0
        self.frame_delays = {
            "IDLE": 15,
            "ATTACK": 10  # 4 frame tấn công hướng xuống
        }

        # Load enemy animations (6 frames IDLE, 4 frames ATTACK)
        self.animations = {
            "IDLE": load_spritesheet("assets/enemy_idle.png", 6, self.sprite_width, self.sprite_height),
        }
        self.current_animation = "IDLE"
        
        # Attack animation từ EnemyAttackAnimator (4 frames)
        self.attack_animator = EnemyAttackAnimator(
            frame_width=self.sprite_width,
            frame_height=self.sprite_height,
            frame_duration=0.12  # ~0.12s mỗi frame cho attack
        )
        
        # Projectile system
        self.projectiles = []
        self.attack_timer = 0
        self.attack_cooldown = 3  # Bắn mỗi 3 giây
        self.last_attack_frame = -1  # Track frame trước đó để tránh bắn nhiều lần

        if boss:
            # Boss version - lớn hơn
            self.sprite_width = 120
            self.sprite_height = 130
            self.animations["IDLE"] = [pygame.transform.scale(frame, (120, 130)) for frame in self.animations["IDLE"]]
            # Re-initialize attack animator với kích thước boss
            self.attack_animator = EnemyAttackAnimator(
                frame_width=120,
                frame_height=130,
                frame_duration=0.12
            )
            # Boss bắn nhanh hơn
            self.attack_cooldown = 2

    def take_damage(self, dmg, game=None):
        old_hp_percent = self.hp / self.max_hp
        self.hp = max(0, self.hp - dmg)
        self.hit_timer = 12
        if self.hp <= 0:
            self.dead = True
            self.death_timer = 60
        
        # Check for boss question milestones
        if self.boss and game and not self.dead:
            new_hp_percent = self.hp / self.max_hp
            for milestone in game.boss_question_milestones:
                if old_hp_percent > milestone >= new_hp_percent and milestone not in game.boss_questions_asked:
                    game.boss_questions_asked.append(milestone)
                    game.trigger_boss_question(milestone)
                    break

    def update(self, dt=None, player=None, game=None):
        if dt is None:
            dt = 1.0 / FPS  # 60 FPS
            
        self.frame += 1
        
        # Cập nhật attack animation
        if self.attack_animator.is_playing:
            self.attack_animator.update(dt)
            
            # Bắn projectile ở frame 2 và 3 (0-indexed: frame 1 và 2)
            current_attack_frame = self.attack_animator.current_frame
            if current_attack_frame in [1, 2] and current_attack_frame != self.last_attack_frame and player:
                self.fire_projectile(player.x, player.y)
            self.last_attack_frame = current_attack_frame
            
            # Reset last_attack_frame khi animation kết thúc
            if not self.attack_animator.is_playing:
                self.last_attack_frame = -1
        
        # Cập nhật idle animation frame (chỉ nếu không attack)
        if not self.attack_animator.is_playing:
            self.frame_timer += 1
            current_delay = self.frame_delays.get("IDLE", 15)
            if self.frame_timer >= current_delay:
                self.frame_timer = 0
                self.current_frame += 1
                idle_frames = len(self.animations.get("IDLE", []))
                if idle_frames > 0 and self.current_frame >= idle_frames:
                    self.current_frame = 0
        
        if self.hit_timer > 0:
            self.hit_timer -= 1
        if self.dead and self.death_timer > 0:
            self.death_timer -= 1
        
        # Cập nhật projectile cooldown (chỉ nếu không chết) - DISABLED: giờ bắn theo attack animation
        # if not self.dead and player:
        #     self.attack_timer += dt
        #     if self.attack_timer >= self.attack_cooldown:
        #         self.fire_projectile(player.x, player.y)
        #         self.attack_timer = 0
        
        # Cập nhật projectiles
        for proj in self.projectiles:
            proj.update(dt)
            
            # Kiểm tra va chạm với player
            if proj.active and proj.get_rect().colliderect(player.get_rect()):
                # Gây sát thương
                actual_damage = int(proj.damage * player.bleed_mult)
                player.hp = max(0, player.hp - actual_damage)
                
                # Hiển thị damage
                game.add_float(player.x, player.y - 30, f"-{actual_damage}", (255, 100, 100), 20)
                
                # Tạo hiệu ứng particles
                game.spawn_particles(player.x, player.y, (255, 100, 100), 8)
                
                # Deactivate projectile
                proj.active = False
                
        self.projectiles = [p for p in self.projectiles if p.active]
        
        # Cập nhật debuffs
        self.update_debuffs()

    def fire_projectile(self, target_x, target_y):
        """Bắn một quả đạn về phía player"""
        projectile = Projectile(
            x=self.x,
            y=self.y,
            target_x=target_x,
            target_y=target_y,
            speed=250,
            damage=self.atk
        )
        self.projectiles.append(projectile)

    def update_debuffs(self):
        """Cập nhật các debuff và áp dụng hiệu ứng."""
        # Burn damage (apply once per second)
        if self.burn_damage > 0:
            self.burn_tick += 1
            if self.burn_tick >= FPS:
                self.burn_tick = 0
                self.hp = max(0, self.hp - self.burn_damage)
                # Hiển thị damage từ burn
        
        # Update timers
        to_remove = []
        for debuff, timer in self.debuff_timers.items():
            self.debuff_timers[debuff] = timer - 1
            if self.debuff_timers[debuff] <= 0:
                to_remove.append(debuff)
        
        for debuff in to_remove:
            del self.debuff_timers[debuff]
            if debuff == "stun":
                self.stunned = False
            elif debuff == "damage_reduction":
                self.damage_reduction = 1.0
            elif debuff == "burn":
                self.burn_damage = 0
                self.burn_tick = 0
            elif debuff == "bleed":
                self.bleed_mult = 1.0

    def get_rect(self):
        idle_frames = self.animations.get("IDLE", [])
        if idle_frames:
            return idle_frames[self.current_frame].get_rect(center=(self.x, self.y))
        return pygame.Rect(self.x - self.sprite_width//2, self.y - self.sprite_height//2, 
                          self.sprite_width, self.sprite_height)

    def draw(self, surf, player=None):
        bob = math.sin(self.frame * 0.05) * 4
        # Logic nháy đỏ khi trúng đòn
        blink = self.hit_timer > 0 and (self.hit_timer // 3) % 2 == 0
        
        if self.dead:
            alpha = int(255 * self.death_timer / 60)
        else:
            alpha = 255

        ex, ey = self.x, int(self.y + bob)

        # Vẽ bóng (Shadow) - Giữ nguyên từ code cũ cho đẹp
        sh = pygame.Surface((50, 14), pygame.SRCALPHA)
        pygame.draw.ellipse(sh, (0, 0, 0, 80), (0, 0, 50, 14))
        surf.blit(sh, (ex - 25, ey + 28))

        # Lấy animation frame hiện tại - từ attack animator nếu đang tấn công, ngược lại từ idle
        if self.attack_animator.is_playing:
            current_sprite = self.attack_animator.get_current_frame()
            if current_sprite is None:
                # Fallback nếu frame không tồn tại
                current_sprite = pygame.Surface((self.sprite_width, self.sprite_height), pygame.SRCALPHA)
            else:
                # Quay attack sprite hướng tới player
                if player:
                    angle = math.atan2(player.y - self.y, player.x - self.x)
                    angle_deg = math.degrees(angle)
                    # Rotate sprite để hướng tới player
                    current_sprite = pygame.transform.rotate(current_sprite, -angle_deg + 90)
                current_sprite = current_sprite.copy()
        else:
            # Sử dụng idle animation
            anim_frames = self.animations.get("IDLE", [])
            if anim_frames:
                current_sprite = anim_frames[self.current_frame].copy()
            else:
                # Fallback nếu không có animation
                current_sprite = pygame.Surface((self.sprite_width, self.sprite_height), pygame.SRCALPHA)
        
        # Hiệu ứng nháy trắng/đỏ khi bị thương
        if blink:
            current_sprite.fill((255, 100, 100, 150), special_flags=pygame.BLEND_RGBA_ADD)
            
        current_sprite.set_alpha(alpha)
        
        # Vẽ sprite (căn giữa)
        sprite_offset_x = current_sprite.get_width() // 2
        sprite_offset_y = current_sprite.get_height() // 2
        surf.blit(current_sprite, (ex - sprite_offset_x, ey - sprite_offset_y))

        # HP bar - Giữ nguyên logic hiển thị máu
        bar_w = 120 if self.boss else 90
        bar_h = 12 if self.boss else 10
        bx, by = ex - bar_w // 2, ey - 70
        pygame.draw.rect(surf, C_HP_BG, (bx, by, bar_w, bar_h), border_radius=5)
        fill_w = int(bar_w * self.hp / self.max_hp)
        if fill_w > 0:
            pygame.draw.rect(surf, C_HP_RED, (bx, by, fill_w, bar_h), border_radius=5)
        if self.boss:
            try:
                font = pygame.font.SysFont("segoeui", 16, bold=True)
            except:
                font = pygame.font.Font(None, 16)
            label = font.render("BOSS", True, (255, 120, 50))
            surf.blit(label, (ex - label.get_width() // 2, by - 24))


# ==================== PLAYER ====================

# ==================== SWORD EFFECT ====================

class SwordEffect:
    def __init__(self, x, y, sprite):
        self.x = x
        self.y = y
        self.sprite = sprite
        self.vx = 0
        self.vy = -18  # Bay lên phía quái
        self.life = 35
        self.max_life = 35
        self.hit = False

    def update(self):
        self.x += self.vx
        self.y += self.vy
        self.vy += 0.3
        self.life -= 1

    def get_rect(self):
        return self.sprite.get_rect(center=(self.x, self.y))

    def draw(self, surf):
        if self.life <= 0:
            return

        alpha = int(255 * self.life / self.max_life)
        sprite = self.sprite.copy()
        sprite.set_alpha(alpha)
        surf.blit(sprite, (int(self.x - sprite.get_width() // 2), int(self.y - sprite.get_height() // 2)))


class Player:
    def __init__(self):
        self.x = SCREEN_W // 2
        self.y = 540
        self.hp = 150
        self.max_hp = 150
        self.gold = 0
        self.level = 1
        self.xp = 0
        self.xp_next = 50
        self.damage_mult = 1.5
        self.combo_bonus = 0.09
        self.lifesteal_percent = 0.25  # 25% of damage dealt

        # Debuffs
        self.stunned = False
        self.damage_reduction = 1.0  # Multiplier for damage dealt (reduced when debuffed)
        self.burn_damage = 0  # Damage per turn from burn
        self.burn_tick = 0
        self.bleed_mult = 1.0  # Multiplier for damage received (increased when bleed)
        self.debuff_timers = {}  # Track debuff durations

        # Kích thước sprite khớp với khung xương (224x278)
        self.sprite_width = 224
        self.sprite_height = 278

        # --- CẤU HÌNH ANIMATION ---
        self.state = "WALK"  # Bắt đầu với animation đi bộ
        self.frame_timer = 0
        self.current_frame = 0
        self.frame_delays = {
            "WALK": 12,
            "ATTACK": 15  # Chậm hơn cho 4 frame
        }

        # Load animations từ spritesheet
        self.animations = {
            "WALK": load_spritesheet("assets/player_walk.png", 4, self.sprite_width, self.sprite_height),
            "ATTACK": [pygame.transform.flip(frame, True, False) for frame in load_spritesheet("assets/player_attack.png", 4, self.sprite_width, self.sprite_height)]  # Lật ngang để hướng chém đúng
        }
        
        # Load sprite cho kiếm khí (projectile)
        raw_sword = pygame.image.load("assets/kiemkhi_proj.png").convert_alpha()
        raw_sword = pygame.transform.rotate(raw_sword, -90)
        raw_sword = pygame.transform.flip(raw_sword, True, False)  # Lật ngược kiếm khí
        self.sword_sprite = pygame.transform.smoothscale(raw_sword, (220, 220))
        
        # Hiệu ứng kiếm khí
        self.sword_effects = []

    def change_state(self, new_state):
        if self.state != new_state and new_state in self.animations:
            self.state = new_state
            self.current_frame = 0
            self.frame_timer = 0

    def update(self):
        # Cập nhật frame animation
        self.frame_timer += 1
        if self.frame_timer >= self.frame_delays.get(self.state, 12):
            self.frame_timer = 0
            old_frame = self.current_frame
            self.current_frame += 1

            # Tạo hiệu ứng kiếm khí ở frame thứ 3 của attack
            if self.state == "ATTACK" and old_frame == 2 and self.current_frame == 3:
                effect_x = self.x
                effect_y = self.y - 60
                self.sword_effects.append(SwordEffect(effect_x, effect_y, self.sword_sprite))

            # Kiểm tra nếu hết chu kỳ animation
            anim_frames = len(self.animations.get(self.state, []))
            if anim_frames > 0 and self.current_frame >= anim_frames:
                if self.state == "ATTACK":
                    self.change_state("WALK")  # Đánh xong thì quay về đi bộ
                else:
                    self.current_frame = 0  # Lặp lại WALK
        
        # Cập nhật hiệu ứng kiếm khí
        self.sword_effects = [effect for effect in self.sword_effects if effect.life > 0]
        for effect in self.sword_effects:
            effect.update()
        
        # Cập nhật debuffs
        self.update_debuffs()

    def update_debuffs(self):
        """Cập nhật các debuff và áp dụng hiệu ứng."""
        # Burn damage (apply once per second)
        if self.burn_damage > 0:
            self.burn_tick += 1
            if self.burn_tick >= FPS:
                self.burn_tick = 0
                self.hp = max(0, self.hp - self.burn_damage)
                # Hiển thị damage từ burn
        
        # Update timers
        to_remove = []
        for debuff, timer in self.debuff_timers.items():
            self.debuff_timers[debuff] = timer - 1
            if self.debuff_timers[debuff] <= 0:
                to_remove.append(debuff)
        
        for debuff in to_remove:
            del self.debuff_timers[debuff]
            if debuff == "stun":
                self.stunned = False
            elif debuff == "damage_reduction":
                self.damage_reduction = 1.0
            elif debuff == "burn":
                self.burn_damage = 0
                self.burn_tick = 0
            elif debuff == "bleed":
                self.bleed_mult = 1.0

    def level_up(self):
        self.level += 1
        self.xp = 0
        self.xp_next = int(self.xp_next * 1.5)
        self.max_hp += 20
        self.hp = min(self.max_hp, self.hp + 35)

    def draw(self, surf):
        # Hiệu ứng nhấp nhô nhẹ nhàng
        bob = math.sin(pygame.time.get_ticks() * 0.005) * 3

        # Lấy frame hiện tại từ bộ animations
        current_anim = self.animations.get(self.state, [])
        if not current_anim:
            return
        
        sprite = current_anim[self.current_frame % len(current_anim)]

        # Vẽ nhân vật với sprite lớn hơn
        px, py = self.x, int(self.y + bob)

        # Hiệu ứng Attack (Tiến về trước một chút rồi đứng yên)
        offset_x = 0
        if self.state == "ATTACK":
            # Tiến về trước trong 2 frame đầu, rồi đứng yên
            if self.current_frame < 2:
                offset_x = self.current_frame * 10  # Tiến 10px mỗi frame

        sprite_rect = sprite.get_rect(center=(px + offset_x, py))
        surf.blit(sprite, sprite_rect)

        # Vẽ hiệu ứng kiếm khí nếu có
        for effect in self.sword_effects:
            effect.draw(surf)
        
    def get_rect(self) -> pygame.Rect:
        """Lấy rect để kiểm tra va chạm"""
        return pygame.Rect(self.x - self.sprite_width//2, self.y - self.sprite_height//2,
                          self.sprite_width, self.sprite_height)


# ==================== SCENE / BACKGROUND ====================

def draw_background(surf, frame, path_offset):
    w, h = SCREEN_W, SCREEN_H

    # Sky gradient
    for y in range(h // 2):
        t = y / (h // 2)
        r = int(5 + t * 15)
        g = int(3 + t * 8)
        b = int(18 + t * 40)
        pygame.draw.line(surf, (r, g, b), (0, y), (w, y))

    # Moon glow
    moon_surf = pygame.Surface((200, 200), pygame.SRCALPHA)
    pygame.draw.circle(moon_surf, (100, 80, 180, 40), (100, 100), 95)
    pygame.draw.circle(moon_surf, (180, 160, 255, 30), (100, 100), 80)
    pygame.draw.circle(moon_surf, (220, 210, 255, 100), (100, 100), 58)
    surf.blit(moon_surf, (w//2 - 100, -30))

    # Castle
    draw_castle(surf, w//2, 165)

    # Perspective path
    top_w = 80
    bot_w = 340
    top_y = 140
    bot_y = h

    path_pts = [
        (w//2 - top_w//2, top_y),
        (w//2 + top_w//2, top_y),
        (w//2 + bot_w//2, bot_y),
        (w//2 - bot_w//2, bot_y),
    ]
    pygame.draw.polygon(surf, C_PATH, path_pts)

    # Path cracks
    for i in range(6):
        t = ((path_offset * 0.003 + i / 6) % 1.0)
        y = int(top_y + (bot_y - top_y) * t)
        hw = (top_w//2 + (bot_w//2 - top_w//2) * t)
        x1 = w//2 - int(hw * 0.3) + random.randint(-2, 2)
        x2 = w//2 + int(hw * 0.2)
        pygame.draw.line(surf, C_PATH2, (x1, y), (x2, y + 12), 1)

    # Tombstones
    draw_tombstone(surf, w*0.15, 340, 0.75)
    draw_tombstone(surf, w*0.8,  320, 0.85)
    draw_tombstone(surf, w*0.08, 480, 1.0)
    draw_tombstone(surf, w*0.88, 450, 0.9)

    # Trees
    draw_dead_tree(surf, 30, 250)
    draw_dead_tree(surf, w - 50, 270)

    # Lanterns
    flicker = math.sin(frame * 0.07) * 0.3 + 0.7
    draw_lantern(surf, w*0.26, 430, flicker)
    draw_lantern(surf, w*0.72, 410, flicker)

    # Bats
    for i in range(3):
        bat_x = w//2 + math.sin(frame * 0.03 + i * 2.1) * 160
        bat_y = 50 + math.cos(frame * 0.04 + i * 1.7) * 25 + i * 20
        draw_bat(surf, int(bat_x), int(bat_y), frame + i * 20)


def draw_castle(surf, cx, base_y):
    # Body
    pygame.draw.rect(surf, (10, 6, 25), (cx - 65, base_y - 70, 130, 90))
    # Side towers
    pygame.draw.rect(surf, (8, 4, 20), (cx - 90, base_y - 105, 38, 115))
    pygame.draw.rect(surf, (8, 4, 20), (cx + 52, base_y - 100, 38, 110))
    # Battlements main
    for i in range(5):
        pygame.draw.rect(surf, (8, 4, 20), (cx - 55 + i*28, base_y - 85, 18, 20))
    # Battlements towers
    for i in range(4):
        pygame.draw.rect(surf, (8, 4, 20), (cx - 94 + i*12, base_y - 118, 8, 15))
        pygame.draw.rect(surf, (8, 4, 20), (cx + 50 + i*12, base_y - 113, 8, 15))
    # Windows (glow)
    win = pygame.Surface((22, 28), pygame.SRCALPHA)
    pygame.draw.rect(win, (220, 160, 40, 140), (0, 0, 22, 28), border_radius=11)
    surf.blit(win, (cx - 11, base_y - 58))
    # Side windows
    w2 = pygame.Surface((14, 18), pygame.SRCALPHA)
    pygame.draw.rect(w2, (200, 140, 30, 100), (0, 0, 14, 18), border_radius=7)
    surf.blit(w2, (cx - 80, base_y - 80))
    surf.blit(w2, (cx + 66, base_y - 76))


def draw_tombstone(surf, x, y, scale):
    w, h = int(28 * scale), int(40 * scale)
    x, y = int(x - w//2), int(y - h)
    pygame.draw.rect(surf, (22, 14, 40), (x, y + h//3, w, h*2//3), border_radius=3)
    pygame.draw.rect(surf, (25, 16, 45), (x, y, w, h//2), border_radius=w//2)
    pygame.draw.rect(surf, (30, 20, 52), (x + 2, y + 2, w - 4, h//2 - 4), border_radius=w//2 - 2)


def draw_dead_tree(surf, x, y):
    pygame.draw.line(surf, (15, 10, 28), (x, y), (x, y - 80), 5)
    pygame.draw.line(surf, (15, 10, 28), (x, y - 50), (x - 30, y - 80), 3)
    pygame.draw.line(surf, (15, 10, 28), (x, y - 40), (x + 25, y - 65), 3)
    pygame.draw.line(surf, (15, 10, 28), (x, y - 65), (x - 20, y - 90), 2)


def draw_lantern(surf, x, y, flicker):
    x, y = int(x), int(y)
    glow = pygame.Surface((60, 60), pygame.SRCALPHA)
    a = int(120 * flicker)
    pygame.draw.circle(glow, (255, 160, 40, a//3), (30, 30), 28)
    pygame.draw.circle(glow, (255, 180, 60, a//2), (30, 30), 18)
    surf.blit(glow, (x - 30, y - 30))
    pygame.draw.circle(surf, (255, int(160*flicker), 40), (x, y), 7)
    pygame.draw.line(surf, (100, 80, 40), (x, y - 7), (x, y - 22), 2)


def draw_bat(surf, x, y, frame):
    wing = math.sin(frame * 0.3) * 8
    pygame.draw.ellipse(surf, (20, 10, 35), (x - 14, y - 3 + int(wing), 12, 8))
    pygame.draw.ellipse(surf, (20, 10, 35), (x + 2, y - 3 - int(wing), 12, 8))
    pygame.draw.circle(surf, (25, 15, 40), (x, y), 5)


# ==================== UI PANELS ====================

def draw_hud(surf, player, level):
    # Top bar
    bar_surf = pygame.Surface((SCREEN_W, 52), pygame.SRCALPHA)
    bar_surf.fill((0, 0, 0, 160))
    surf.blit(bar_surf, (0, 0))

    # Level badge
    draw_rounded_rect(surf, (70, 35, 130), pygame.Rect(8, 8, 58, 24), 6)
    draw_text_center(surf, f"Lv.{player.level}", font_sm, C_LAVENDER, 37, 20)

    # HP bar
    bar_x, bar_y, bar_w, bar_h = 75, 12, 220, 16
    pygame.draw.rect(surf, (40, 20, 60), (bar_x, bar_y, bar_w, bar_h), border_radius=8)
    hp_w = int(bar_w * player.hp / player.max_hp)
    if hp_w > 0:
        pygame.draw.rect(surf, C_HP_RED, (bar_x, bar_y, hp_w, bar_h), border_radius=8)
    pygame.draw.rect(surf, (180, 80, 80), (bar_x, bar_y, bar_w, bar_h), 1, border_radius=8)
    draw_text(surf, f"HP {player.hp}/{player.max_hp}", font_xs, C_WHITE, bar_x + 4, bar_y + 2, shadow=False)

    # XP bar
    xp_x, xp_y, xp_w, xp_h = 75, 30, 220, 8
    pygame.draw.rect(surf, (20, 10, 40), (xp_x, xp_y, xp_w, xp_h), border_radius=4)
    xx = int(xp_w * min(player.xp / player.xp_next, 1.0))
    if xx > 0:
        pygame.draw.rect(surf, C_PURPLE, (xp_x, xp_y, xx, xp_h), border_radius=4)

    # Gold
    gold_txt = font_md.render(f"💰 {player.gold}", True, C_GOLD)
    surf.blit(gold_txt, (SCREEN_W - gold_txt.get_width() - 10, 12))


def draw_dice_panel(surf, dice_list, hand_name, damage, rerolls, max_rerolls, combo_count=0, combo_effect=None):
    panel_y = SCREEN_H - 210
    panel_h = 210

    # Panel background
    panel = pygame.Surface((SCREEN_W, panel_h), pygame.SRCALPHA)
    panel.fill((10, 5, 25, 220))
    surf.blit(panel, (0, panel_y))
    pygame.draw.line(surf, C_PURPLE, (0, panel_y), (SCREEN_W, panel_y), 2)

    # Hand name bar
    hand_color = HAND_COLORS.get(hand_name, C_WHITE)
    hw_surf = pygame.Surface((SCREEN_W, 28), pygame.SRCALPHA)
    pygame.draw.rect(hw_surf, (*hand_color, 30), (0, 0, SCREEN_W, 28))
    surf.blit(hw_surf, (0, panel_y))

    draw_text_center(surf, hand_name, font_hand, hand_color, SCREEN_W // 2, panel_y + 14)
    if combo_count >= 3:
        combo_txt = font_sm.render(f"{combo_count}x COMBO!", True, C_GOLD)
        surf.blit(combo_txt, (10, panel_y + 8))
        if combo_effect:
            effect_txt = font_xs.render(combo_effect.replace("_", " "), True, C_CYAN)
            surf.blit(effect_txt, (10, panel_y + 28))
    if damage > 0:
        dmg_txt = font_hand.render(f"💥 {damage} DMG", True, C_RED)
        surf.blit(dmg_txt, (SCREEN_W - dmg_txt.get_width() - 10, panel_y + 5))

    # Dice row
    die_y = panel_y + 70
    padding = 35  # Increased margin from edges to ensure all dice fit
    available_width = SCREEN_W - 2 * padding
    spacing = available_width / len(dice_list)
    for i, die in enumerate(dice_list):
        die.x = padding + spacing * (i + 0.5)
        die.y = die_y
        die.draw(surf)

    # Reroll indicator
    rr_x = 14
    rr_y = panel_y + 140
    draw_text(surf, "Tung lại:", font_sm, C_GRAY, rr_x, rr_y)
    for i in range(max_rerolls):
        c = C_PURPLE if i < rerolls else (40, 25, 60)
        pygame.draw.circle(surf, c, (rr_x + 90 + i * 24, rr_y + 9), 9)
        pygame.draw.circle(surf, C_LAVENDER if i < rerolls else (60, 40, 80), (rr_x + 90 + i * 24, rr_y + 9), 9, 1)

    # Tip text
    tip = font_xs.render("Click xúc xắc để khóa | SPACE: tung lại | ENTER: tấn công", True, (100, 80, 140))
    surf.blit(tip, (SCREEN_W//2 - tip.get_width()//2, panel_y + 162))

    # Button hints
    draw_rounded_rect(surf, (60, 30, 110), pygame.Rect(10, panel_y + 180, 110, 24), 6)
    draw_text_center(surf, "[SPACE] Tung lại", font_xs, C_LAVENDER, 65, panel_y + 192)
    draw_rounded_rect(surf, (120, 30, 30), pygame.Rect(SCREEN_W - 130, panel_y + 180, 120, 24), 6)
    draw_text_center(surf, "[ENTER] Tấn công", font_xs, (255, 140, 140), SCREEN_W - 70, panel_y + 192)


def draw_damage_number(surf, x, y, value):
    """Draws a large damage number above enemy"""
    txt = font_lg.render(f"{value}", True, C_GOLD)
    surf.blit(txt, (x - txt.get_width()//2, y))


def draw_level_up_banner(surf, level):
    overlay = pygame.Surface((SCREEN_W, 80), pygame.SRCALPHA)
    overlay.fill((80, 40, 200, 180))
    surf.blit(overlay, (0, SCREEN_H//2 - 40))
    draw_text_center(surf, f"⬆ LEVEL UP! Lv.{level}", font_lg, C_GOLD, SCREEN_W//2, SCREEN_H//2)
    draw_text_center(surf, "+20 HP Max   +30 HP hồi phục", font_sm, C_WHITE, SCREEN_W//2, SCREEN_H//2 + 26)


def draw_level_up_menu(surf, level, options, selected_index):
    overlay = pygame.Surface((SCREEN_W, SCREEN_H), pygame.SRCALPHA)
    overlay.fill((0, 0, 0, 180))
    surf.blit(overlay, (0, 0))

    title = font_lg.render(f"LEVEL UP! Lv.{level}", True, C_GOLD)
    surf.blit(title, (SCREEN_W//2 - title.get_width()//2, 80))

    subtitle = font_sm.render("Chọn thuộc tính nâng cấp", True, C_WHITE)
    surf.blit(subtitle, (SCREEN_W//2 - subtitle.get_width()//2, 125))

    box_w = 340
    box_h = 90
    start_y = 170
    for i, option in enumerate(options):
        rect = pygame.Rect(SCREEN_W//2 - box_w//2, start_y + i * (box_h + 14), box_w, box_h)
        box_color = (80, 45, 140) if i == selected_index else (40, 20, 70)
        border_color = C_GOLD if i == selected_index else (90, 60, 140)
        draw_rounded_rect(surf, box_color, rect, 16)
        pygame.draw.rect(surf, border_color, rect, 2, border_radius=16)

        txt = font_md.render(option["title"], True, C_WHITE)
        desc = font_xs.render(option["desc"], True, (220, 220, 240))
        surf.blit(txt, (rect.x + 18, rect.y + 18))
        surf.blit(desc, (rect.x + 18, rect.y + 46))

    hint = font_xs.render("Dùng ↑/↓ chọn và Enter xác nhận", True, C_LAVENDER)
    surf.blit(hint, (SCREEN_W//2 - hint.get_width()//2, start_y + len(options) * (box_h + 14) + 10))


def draw_game_over(surf):
    overlay = pygame.Surface((SCREEN_W, SCREEN_H), pygame.SRCALPHA)
    overlay.fill((0, 0, 0, 180))
    surf.blit(overlay, (0, 0))
    draw_text_center(surf, "GAME OVER", font_lg, C_RED, SCREEN_W//2, SCREEN_H//2 - 30)
    draw_text_center(surf, "Nhấn R để chơi lại", font_md, C_WHITE, SCREEN_W//2, SCREEN_H//2 + 20)


def draw_menu(surf, selected_index, has_save, game=None):
    # Background
    surf.fill(C_BG)
    
    # Title
    title = font_lg.render("⚔️ DICE ROGUELITE", True, C_GOLD)
    surf.blit(title, (SCREEN_W//2 - title.get_width()//2, 100))
    
    # Menu options
    options = ["Chơi mới", "Chơi tiếp" if has_save else "Chơi tiếp (không có save)", "Tạo câu hỏi từ file", "Thoát"]
    start_y = 250
    box_w = 300
    box_h = 50
    
    for i, option in enumerate(options):
        rect = pygame.Rect(SCREEN_W//2 - box_w//2, start_y + i * (box_h + 20), box_w, box_h)
        box_color = (80, 45, 140) if i == selected_index else (40, 20, 70)
        border_color = C_GOLD if i == selected_index else (90, 60, 140)
        draw_rounded_rect(surf, box_color, rect, 12)
        pygame.draw.rect(surf, border_color, rect, 2, border_radius=12)
        
        txt = font_md.render(option, True, C_WHITE)
        surf.blit(txt, (rect.x + (box_w - txt.get_width())//2, rect.y + (box_h - txt.get_height())//2))
    
    # Import status message
    if game and game.import_status:
        status_color = C_GREEN if "Thành công" in game.import_status else C_RED
        status_text = font_sm.render(game.import_status, True, status_color)
        surf.blit(status_text, (SCREEN_W//2 - status_text.get_width()//2, start_y + len(options) * (box_h + 20) - 40))


def draw_questions(surf, game):
    # Background
    surf.fill(C_BG)
    
    # Title
    title = font_lg.render("HƯỚNG DẪN", True, C_GOLD)
    surf.blit(title, (SCREEN_W//2 - title.get_width()//2, 50))
    
    if game.all_questions:
        # Hiển thị danh sách câu hỏi
        subtitle = font_md.render("CÂU HỎI TRONG GAME", True, C_LAVENDER)
        surf.blit(subtitle, (SCREEN_W//2 - subtitle.get_width()//2, 100))
        
        start_y = 150
        displayed = 0
        for q_data in game.all_questions:
            if displayed >= 8:  # Giới hạn hiển thị
                if len(game.all_questions) > 8:
                    more = font_xs.render(f"...và {len(game.all_questions) - 8} câu khác", True, C_GRAY)
                    surf.blit(more, (SCREEN_W//2 - more.get_width()//2, start_y + displayed * 30))
                break
            text = f"{q_data['id']}. {q_data['text']}"
            txt = font_sm.render(text, True, C_WHITE)
            surf.blit(txt, (30, start_y + displayed * 30))
            displayed += 1
    else:
        # Hiển thị hướng dẫn chơi game
        guide_title = font_md.render("HƯỚNG DẪN CHƠI GAME", True, C_LAVENDER)
        surf.blit(guide_title, (SCREEN_W//2 - guide_title.get_width()//2, 100))
        
        guide = [
            "SPACE: Tung lại xúc xắc",
            "ENTER: Tấn công kẻ thù", 
            "Click: Khóa/mở xúc xắc",
            "R: Chơi lại khi thua",
            "",
            "Sử dụng: Tạo câu hỏi từ file",
            "để thêm câu hỏi custom vào game"
        ]
        start_y = 150
        for i, text in enumerate(guide):
            if text:
                color = C_CYAN if i > 3 else C_WHITE
                txt = font_md.render(text, True, color)
            else:
                continue
            surf.blit(txt, (SCREEN_W//2 - txt.get_width()//2, start_y + i * 40))
    
    # Back hint
    hint = font_sm.render("Nhấn ESC để quay lại menu", True, C_LAVENDER)
    surf.blit(hint, (SCREEN_W//2 - hint.get_width()//2, SCREEN_H - 60))




def draw_boss_question(surf, game, selected_option):
    """Vẽ câu hỏi khi fight boss."""
    if not game.boss_question:
        return
    
    question = game.boss_question
    # Nếu chưa có lựa chọn, mặc định là 0
    if selected_option < 0:
        selected_option = 0
    
    # Semi-transparent background
    s = pygame.Surface((SCREEN_W, 350), pygame.SRCALPHA)
    pygame.draw.rect(s, (20, 10, 40, 200), s.get_rect(), border_radius=10)
    surf.blit(s, (0, SCREEN_H - 350))
    
    # Title
    title = font_md.render("⚡ CÂU HỎI ⚡", True, C_GOLD)
    surf.blit(title, (SCREEN_W//2 - title.get_width()//2, SCREEN_H - 330))
    
    # Debuff hint
    if game.current_debuff_name:
        debuff_descriptions = {
            "Choáng": "Không thể tấn công",
            "Giảm Sát Thương": "Sát thương giảm 50%",
            "Thiêu Đốt": "Mỗi giây mất 5 HP",
            "Trọng Thương": "Nhận 50% sát thương thêm"
        }
        desc = debuff_descriptions.get(game.current_debuff_name, "")
        hint_title = font_md.render(f"⚠️ {game.current_debuff_name}", True, C_RED)
        hint_desc = font_xs.render(desc, True, C_ORANGE)
        surf.blit(hint_title, (SCREEN_W//2 - hint_title.get_width()//2, SCREEN_H - 310))
        surf.blit(hint_desc, (SCREEN_W//2 - hint_desc.get_width()//2, SCREEN_H - 285))
        start_y = SCREEN_H - 250
    else:
        start_y = SCREEN_H - 300

    # Question text (word wrap)
    question_text = question["text"]
    words = question_text.split()
    lines = []
    current_line = ""
    for word in words:
        test_line = current_line + word + " "
        txt = font_md.render(test_line, True, C_WHITE)
        if txt.get_width() > SCREEN_W - 60:
            if current_line:
                lines.append(current_line)
            current_line = word + " "
        else:
            current_line = test_line
    if current_line:
        lines.append(current_line)
    
    start_y = SCREEN_H - 300
    for line in lines:
        txt = font_md.render(line, True, C_WHITE)
        surf.blit(txt, (SCREEN_W//2 - txt.get_width()//2, start_y))
        start_y += 35
    
    # Options
    start_y = SCREEN_H - 150
    for i, option in enumerate(question["options"]):
        color = C_GOLD if i == selected_option else C_WHITE
        marker = "→ " if i == selected_option else "  "
        txt = font_sm.render(f"{marker}{chr(65+i)}: {option}", True, color)
        surf.blit(txt, (40, start_y + i * 30))
    
    # Hint
    hint = font_xs.render("Mũi tên ↑↓ chọn, ENTER xác nhận", True, C_LAVENDER)
    surf.blit(hint, (SCREEN_W//2 - hint.get_width()//2, SCREEN_H - 20))


# ==================== MAIN GAME ====================

class GameState(Enum):
    MENU      = "menu"
    ROLLING   = "rolling"
    ATTACKING = "attacking"
    ENEMY_ATK = "enemy_atk"
    LEVEL_UP  = "level_up"
    GAME_OVER = "game_over"


class Game:
    def __init__(self):
        self.state = GameState.MENU
        self.menu_selected = 0
        self.has_save = False
        # Initialize basic attributes
        self.frame = 0
        self.path_offset = 0
        self.particles = []
        self.float_texts = []
        self.player = None
        self.enemy = None
        self.dice = []
        self.show_questions = False
        self.shake_timer = 0
        self.map_themes = ["lava", "icemap", "naturemap"]
        self.current_map_index = 0
        self.kills_until_boss = 0
        self.boss_active = False
        
        # --- QUESTIONS SYSTEM ---
        self.all_questions = load_questions()
        self.boss_question = None
        self.boss_question_answered = False
        self.answer_selected = -1
        self.boss_question_milestones = [0.8, 0.6, 0.4, 0.2, 0.01]  # HP percentages for questions
        self.boss_questions_asked = []  # Track which milestones have been asked
        self.boss_wrong_answers = 0
        self.current_debuff_type = None
        self.current_debuff_name = None
        self.import_status = ""
        self.import_status_timer = 0
        
        # Other attributes will be set in reset()

        # --- MAP SYSTEM ---
        self.map_manager = MapManager(SCREEN_W, SCREEN_H)
        self.map_manager.generate_map(
            cols=30, rows=20,
            theme=self.map_themes[self.current_map_index],
            seed=0,  # Fixed seed cho menu
            static_background=True,
            background_image_path=self.get_map_image_path(self.map_themes[self.current_map_index]),
            tile_w=16,
            scale=1.0
        )

    def reset(self):
        self.player    = Player()
        self.enemy     = Enemy(level=1)
        self.dice      = [Die(i) for i in range(5)]
        self.rerolls   = 3
        self.max_rerolls = 3
        self.current_map_index = 0
        self.kills_until_boss = 0
        self.boss_active = False
        self.particles : list[Particle]    = []
        self.float_texts: list[FloatingText] = []
        self.state     = GameState.ROLLING
        self.path_offset = 0
        self.frame     = 0
        self.hand_name = "—"
        self.damage    = 0
        self.combo_count = 0
        self.combo_value = None
        self.combo_effect = None
        self.level_up_options = []
        self.selected_upgrade = 0
        self.enemy_atk_timer = 0
        self.level_up_timer  = 0
        self.level_count = 1
        self.shake_timer = 0   # screen shake frames
        self.pending_attack = None
        self.menu_selected = 0
        self.has_save = False  # TODO: implement save/load later
        self.boss_active = False
        
        # --- RESET QUESTIONS ---
        self.boss_question = None
        self.boss_question_answered = False
        self.answer_selected = -1
        self.boss_questions_asked = []
        self.boss_wrong_answers = 0
        self.current_debuff_type = None
        self.current_debuff_name = None

        # --- MAP SYSTEM ---
        self.map_manager = MapManager(SCREEN_W, SCREEN_H)
        self.map_manager.generate_map(
            cols=30, rows=20,
            theme=self.map_themes[self.current_map_index],
            seed=self.level_count,  # Different map each wave
            static_background=True,
            background_image_path=self.get_map_image_path(self.map_themes[self.current_map_index]),
            tile_w=16,
            scale=1.0
        )

        # Unlock first 2 dice
        self.dice[0].roll(stagger_frames=0)
        self.dice[1].roll(stagger_frames=8)
        self.dice[2].roll(stagger_frames=16)
        self._eval()

    def _eval(self):
        vals = [d.value for d in self.dice if not d.empty]
        self.hand_name, base, mult, self.combo_count, self.combo_value = eval_hand(vals)
        combo_bonus = 1.0 + max(0, self.combo_count - 2) * self.player.combo_bonus
        self.damage = int(base * mult * combo_bonus * self.player.damage_mult) if vals else 0
        self.combo_effect = None
        if self.combo_count >= 3:
            if self.combo_value == 3:
                self.combo_effect = "STUN"
            elif self.combo_value == 4:
                self.combo_effect = "HEAL" if self.combo_count == 3 else "HEAL_BIG"
            elif self.combo_value == 6:
                self.combo_effect = "CRIT"
            else:
                self.combo_effect = "COMBO"
        combo_color = None
        if self.combo_count == 3:
            combo_color = (255, 185, 80)
        elif self.combo_count == 4:
            combo_color = (255, 125, 50)
        elif self.combo_count >= 5:
            combo_color = (255, 110, 220)
        for d in self.dice:
            d.combo_glow = (not d.empty and self.combo_count >= 3 and d.value == self.combo_value)
            d.combo_color = combo_color if d.combo_glow else None

    def spawn_particles(self, x, y, color, n=10):
        for _ in range(n):
            self.particles.append(Particle(x, y, color))

    def add_float(self, x, y, text, color, size=22):
        self.float_texts.append(FloatingText(x, y, text, color, size))

    def handle_reroll(self):
        if self.state != GameState.ROLLING:
            return
        if self.rerolls <= 0:
            return
        any_unlocked = any(not d.empty and not d.locked for d in self.dice)
        if not any_unlocked:
            return

        stagger = 0
        # Unlock a new slot every reroll used (after first reroll)
        used = self.max_rerolls - self.rerolls
        slot_to_unlock = 2 + used
        if used > 0 and slot_to_unlock < 5 and self.dice[slot_to_unlock].empty:
            self.dice[slot_to_unlock].roll(stagger_frames=0)
        else:
            # Re-roll all unlocked dice with stagger offset
            for d in self.dice:
                if not d.empty and not d.locked:
                    d.roll(stagger_frames=stagger)
                    stagger += 5

        self.rerolls -= 1
        self._eval()

    def handle_attack(self):
        if self.state != GameState.ROLLING:
            return
        if self.damage <= 0:
            return
        self.state = GameState.ATTACKING
        self.player.change_state("ATTACK")  # Kích hoạt animation tấn công
        if self.combo_count >= 3:
            self.add_float(self.player.x, self.player.y - 30, f"{self.combo_count}x COMBO!", C_GOLD, 24)
            self.spawn_particles(self.player.x, self.player.y, C_GOLD, 12 + self.combo_count * 5)
            self.spawn_particles(self.enemy.x, self.enemy.y, C_CYAN, 8 + self.combo_count * 4)

        extra_text = None
        heal_amount = 0
        if self.combo_effect == "STUN":
            self.enemy.stunned = True
            self.enemy.debuff_timers["stun"] = 120  # 2 giây
            extra_text = "STUNNED"
        elif self.combo_effect == "HEAL":
            heal_amount = 15
            extra_text = f"HEAL +{heal_amount}"
        elif self.combo_effect == "HEAL_BIG":
            heal_amount = 30
            extra_text = f"HEAL +{heal_amount}"
        crit_bonus = 0
        if self.combo_effect == "CRIT":
            crit_bonus = int(self.damage * 0.25)
            extra_text = f"CRIT +{crit_bonus}"
        if heal_amount > 0:
            self.player.hp = min(self.player.max_hp, self.player.hp + heal_amount)
            self.add_float(self.player.x, self.player.y - 40, f"+{heal_amount} HP", C_WHITE, 20)
        if extra_text:
            self.add_float(self.enemy.x, self.enemy.y - 60, extra_text, C_CYAN, 22)

        self.pending_attack = {
            "damage": self.damage,
            "gold": max(1, self.damage // 10),
            "combo_effect": self.combo_effect,
            "extra_text": extra_text,
            "heal_amount": heal_amount,
            "crit_bonus": crit_bonus,
            "resolved": False,
        }
        # Reset dice after small delay whether hit or not
        pygame.time.set_timer(pygame.USEREVENT + 3, 500, loops=1)

    def on_enemy_dead(self):
        boss_defeated = hasattr(self.enemy, 'boss') and self.enemy.boss
        if boss_defeated:
            self.current_map_index = (self.current_map_index + 1) % len(self.map_themes)
            self.kills_until_boss = 0
            self.boss_active = False
            self.level_count = 1
            self.player.hp = min(self.player.max_hp, self.player.hp + 30)
            self.map_manager.generate_map(
                cols=30, rows=20,
                theme=self.map_themes[self.current_map_index],
                seed=self.level_count,
                static_background=True,
                background_image_path=self.get_map_image_path(self.map_themes[self.current_map_index]),
                tile_w=16,
                scale=1.0
            )
            self.add_float(SCREEN_W // 2, 80, f"Entering {self.map_themes[self.current_map_index].upper()}", C_CYAN, 24)
            self._spawn_new_enemy()
            return

        self.level_count += 1
        self.kills_until_boss += 1
        self.player.hp = min(self.player.max_hp, self.player.hp + 15)
        if self.kills_until_boss >= 4:
            self.kills_until_boss = 0
            self.boss_active = True
            self._spawn_boss()
            return

        # Check player level up
        if self.player.xp >= self.player.xp_next:
            self.player.level_up()
            self.state = GameState.LEVEL_UP
            self.level_up_options = self.build_level_up_options()
            self.selected_upgrade = 0
        else:
            self._spawn_new_enemy()

    def check_sword_collisions(self):
        if not self.pending_attack or self.pending_attack.get("resolved"):
            return
        enemy_rect = self.enemy.get_rect()
        for effect in self.player.sword_effects:
            if effect.life <= 0 or effect.hit:
                continue
            if effect.get_rect().colliderect(enemy_rect):
                effect.hit = True
                effect.life = 0
                self.resolve_pending_attack()
                break

    def resolve_pending_attack(self):
        if not self.pending_attack or self.pending_attack.get("resolved"):
            return
        self.pending_attack["resolved"] = True
        dmg = self.pending_attack["damage"]
        crit_bonus = self.pending_attack.get("crit_bonus", 0)
        
        # Áp dụng damage reduction của enemy
        total_damage = (dmg + crit_bonus) * self.enemy.damage_reduction
        dmg = int(total_damage - crit_bonus)  # Tách lại để hiển thị
        crit_bonus = int(crit_bonus * self.enemy.damage_reduction)
        
        if crit_bonus > 0:
            self.enemy.take_damage(crit_bonus, self)
            self.spawn_particles(self.enemy.x, self.enemy.y, (255, 100, 100), 8)
        self.enemy.take_damage(dmg, self)
        self.spawn_particles(self.enemy.x, self.enemy.y, C_RED, 12)
        self.add_float(self.enemy.x, self.enemy.y - 40, f"-{dmg}", C_RED, 26)
        
        # Lifesteal: hồi máu dựa trên sát thương
        heal_from_damage = int(total_damage * self.player.lifesteal_percent)
        if heal_from_damage > 0:
            self.player.hp = min(self.player.max_hp, self.player.hp + heal_from_damage)
            self.add_float(self.player.x, self.player.y - 30, f"+{heal_from_damage}💧", (100, 200, 255), 18)
            self.spawn_particles(self.player.x, self.player.y, (100, 150, 255), 6)
        
        g = self.pending_attack["gold"]
        self.player.gold += g
        self.add_float(self.enemy.x + 30, self.enemy.y - 10, f"+{g}💰", C_GOLD, 16)
        if self.pending_attack.get("combo_effect") == "STUN":
            self.enemy_stunned = True
        if self.enemy.dead:
            xp = 25 + self.level_count * 8
            self.player.xp += xp
            self.add_float(self.enemy.x, self.enemy.y - 60, f"+{xp} XP", C_CYAN, 18)
            self.spawn_particles(self.enemy.x, self.enemy.y, C_GOLD, 20)
            pygame.time.set_timer(pygame.USEREVENT + 1, 1500, loops=1)
        else:
            pygame.time.set_timer(pygame.USEREVENT + 2, 600, loops=1)

    def build_level_up_options(self):
        choices = [
            {
                "title": "Max HP +20",
                "desc": "Tăng HP tối đa và hồi 20 HP",
                "apply": lambda: self.apply_max_hp_upgrade()
            },
            {
                "title": "Damage +10%",
                "desc": "Tăng sát thương của dice",
                "apply": lambda: self.apply_damage_upgrade()
            },
            {
                "title": "Extra Reroll",
                "desc": "Tăng 1 lượt tung lại",
                "apply": lambda: self.apply_reroll_upgrade()
            },
            {
                "title": "Combo +2%",
                "desc": "Tăng hiệu ứng combo",
                "apply": lambda: self.apply_combo_upgrade()
            }
        ]
        return random.sample(choices, 3)

    def apply_level_up_option(self, index):
        if index < 0 or index >= len(self.level_up_options):
            return
        option = self.level_up_options[index]
        option["apply"]()
        self.add_float(self.player.x, self.player.y - 40, option["title"], C_GOLD, 20)
        self.state = GameState.ROLLING
        self._spawn_new_enemy()

    def apply_max_hp_upgrade(self):
        self.player.max_hp += 20
        self.player.hp = min(self.player.max_hp, self.player.hp + 20)

    def apply_damage_upgrade(self):
        self.player.damage_mult += 0.10

    def apply_reroll_upgrade(self):
        self.max_rerolls += 1
        self.rerolls += 1

    def apply_combo_upgrade(self):
        self.player.combo_bonus += 0.02

    def _spawn_new_enemy(self):
        self.enemy = Enemy(level=self.level_count)
        self.state = GameState.ROLLING

    def _spawn_boss(self):
        self.enemy = Enemy(level=self.level_count, boss=True)
        self.state = GameState.ROLLING
        self.add_float(self.enemy.x, self.enemy.y - 80, "BOSS APPROACHING!", C_RED, 28)
        self.spawn_particles(self.enemy.x, self.enemy.y, (200, 60, 255), 20)
        self.boss_questions_asked = []
        self.boss_question = None
        self.boss_question_answered = False
        self.answer_selected = -1
        self.current_debuff_type = None
        self.current_debuff_name = None
        self.current_milestone = None
        self.boss_wrong_answers = 0

    def trigger_boss_question(self, milestone):
        """Đặt câu hỏi boss tại mốc máu nhất định."""
        if not self.all_questions:
            return
        
        # Chọn câu hỏi ngẫu nhiên
        question = random.choice(self.all_questions)
        self.boss_question = question
        self.boss_question_answered = False
        self.answer_selected = -1
        
        # Chọn debuff ngẫu nhiên cho câu hỏi này
        debuff_map = {
            "stun": "Choáng",
            "damage_reduction": "Giảm Sát Thương",
            "burn": "Thiêu Đốt",
            "bleed": "Trọng Thương"
        }
        self.current_debuff_type = random.choice(list(debuff_map.keys()))
        self.current_debuff_name = debuff_map[self.current_debuff_type]
        self.current_milestone = milestone
        
        # Pause game state để chờ câu trả lời
        self.state = GameState.ROLLING  # Keep rolling but show question

    def submit_boss_answer(self):
        """Xử lý submit đáp án cho câu hỏi boss."""
        if not self.boss_question or self.answer_selected < 0:
            return
        
        self.boss_question_answered = True
        correct = self.answer_selected == self.boss_question["correct"]
        
        debuff_type = self.current_debuff_type or "stun"
        debuff_name = self.current_debuff_name or "Choáng"
        
        debuff_descriptions = {
            "Choáng": "Không thể tấn công",
            "Giảm Sát Thương": "Sát thương giảm 50%",
            "Thiêu Đốt": "Mỗi giây mất 5 HP",
            "Trọng Thương": "Nhận 50% sát thương thêm"
        }
        
        if correct:
            # Đúng: debuff boss và nhận phần thưởng ngẫu nhiên
            self.apply_debuff(self.enemy, debuff_type, 180)  # 3 giây tại 60 FPS
            reward_text = self.grant_random_question_reward()
            msg = f"✓ Boss bị {debuff_name}! {reward_text}"
            color = C_GREEN
        else:
            self.boss_wrong_answers += 1
            if self.boss_wrong_answers >= 3:
                self.handle_boss_question_failure()
                return

            # Sai: debuff player
            self.apply_debuff(self.player, debuff_type, 180)
            msg = f"✗ Bạn bị {debuff_name}: {debuff_descriptions.get(debuff_name, '')}!"
            color = C_RED
        
        # Hiển thị kết quả
        self.add_float(SCREEN_W//2, SCREEN_H//2 - 50, msg, color, 24)
        self.spawn_particles(SCREEN_W//2, SCREEN_H//2, color, 15)
        
        # Reset question
        self.boss_question = None
        self.current_milestone = None
        self.current_debuff_type = None
        self.current_debuff_name = None

    def apply_debuff(self, target, debuff_type, duration):
        """Áp dụng debuff cho target (player hoặc enemy)."""
        target.debuff_timers[debuff_type] = duration
        
        if debuff_type == "stun":
            target.stunned = True
        elif debuff_type == "damage_reduction":
            target.damage_reduction = 0.5  # Giảm 50% damage
        elif debuff_type == "burn":
            target.burn_damage = 5  # 5 damage per turn
        elif debuff_type == "bleed":
            target.bleed_mult = 1.5  # Nhận 50% damage hơn

    def grant_random_question_reward(self):
        """Cấp phần thưởng ngẫu nhiên khi trả lời đúng câu hỏi boss."""
        def heal_reward():
            heal_amount = 25
            self.player.hp = min(self.player.max_hp, self.player.hp + heal_amount)
            return f"Bạn hồi {heal_amount} HP!"

        def reroll_reward():
            self.max_rerolls += 1
            self.rerolls += 1
            return "Bạn nhận thêm 1 lượt reroll!"

        def damage_reward():
            self.player.damage_mult += 0.10
            return "Sát thương tăng thêm 10%!"

        def gold_reward():
            gold_amount = 20
            self.player.gold += gold_amount
            return f"Bạn nhận thêm {gold_amount} vàng!"

        rewards = [heal_reward, reroll_reward, damage_reward, gold_reward]
        reward_fn = random.choice(rewards)
        return reward_fn()

    def handle_boss_question_failure(self):
        """Xử lý khi trả lời sai quá 2 câu boss."""
        self.boss_active = False
        self.level_count = 1
        self.kills_until_boss = 0
        self.boss_questions_asked = []
        self.boss_question = None
        self.boss_question_answered = False
        self.answer_selected = -1
        self.current_debuff_type = None
        self.current_debuff_name = None
        self.current_milestone = None
        self.boss_wrong_answers = 0

        self.player.hp = max(1, self.player.hp)

        self.map_manager.generate_map(
            cols=30, rows=20,
            theme=self.map_themes[self.current_map_index],
            seed=self.level_count,
            static_background=True,
            background_image_path=self.get_map_image_path(self.map_themes[self.current_map_index]),
            tile_w=16,
            scale=1.0
        )
        self.enemy = Enemy(level=1)
        self.state = GameState.ROLLING

        self.add_float(SCREEN_W//2, SCREEN_H//2 - 40, "Boss đã kết liễu! Quay về Wave 1", C_RED, 20)
        self.spawn_particles(self.player.x, self.player.y, C_RED, 20)

    def get_map_image_path(self, theme: str) -> str:
        mapping = {
            "lava": "assets/Firemap.png",
            "icemap": "assets/Icemap.png",
            "naturemap": "assets/Naturemap.png",
        }
        return mapping.get(theme, "assets/Firemap.png")

    def on_enemy_attack(self):
        if self.enemy.dead:
            return
        if self.enemy.stunned:
            self.enemy.stunned = False
            self.add_float(self.enemy.x, self.enemy.y - 40, "STUN BLOCKED", C_CYAN, 22)
            self.spawn_particles(self.enemy.x, self.enemy.y, (120, 220, 255), 12)
            return
        # Trigger attack animation
        self.enemy.attack_animator.play()
        self.enemy.last_attack_frame = -1  # Reset frame tracking
        atk = self.enemy.atk + random.randint(0, 5)
        # Áp dụng bleed multiplier của player
        actual_damage = int(atk * self.player.bleed_mult)
        self.player.hp = max(0, self.player.hp - actual_damage)
        self.add_float(self.player.x, self.player.y - 30, f"-{actual_damage}", (255, 100, 200), 20)
        self.spawn_particles(self.player.x, self.player.y, (200, 50, 150), 6)
        if self.player.hp <= 0:
            self.state = GameState.GAME_OVER

    def reset_dice(self):
        for d in self.dice:
            d.locked = False
            d.empty = True
            d.phase = Die.PHASE_IDLE
        self.dice[0].roll(stagger_frames=0)
        self.dice[1].roll(stagger_frames=6)
        self.dice[2].roll(stagger_frames=12)
        self.rerolls = self.max_rerolls
        self._eval()
        if self.state == GameState.ATTACKING:
            self.state = GameState.ROLLING

    def toggle_lock(self, die_index):
        d = self.dice[die_index]
        if d.empty:
            return
        d.locked = not d.locked

    def update(self):
        self.frame += 1
        self.path_offset += 1
        
        # Calculate delta time
        dt = 1.0 / FPS
        
        # Cập nhật map (camera follow)
        if self.player:
            self.map_manager.update(self.player.x, self.player.y)
            if self.map_manager.editor:
                self.map_manager.editor.update(pygame.mouse.get_pos())
        
        # Update import status timer
        if self.import_status_timer > 0:
            self.import_status_timer -= 1
        else:
            self.import_status = ""
        
        if self.state == GameState.MENU:
            return  # Don't update game objects in menu
        
        self.player.update()
        self.enemy.update(dt, self.player, self)

        # Sword projectile collision
        self.check_sword_collisions()

        # Particles
        for p in self.particles:
            p.update()
        self.particles = [p for p in self.particles if p.life > 0]

        for ft in self.float_texts:
            ft.update()
        self.float_texts = [ft for ft in self.float_texts if ft.life > 0]

        # Update dice + spawn impact particles on landing
        for d in self.dice:
            prev_phase = d.phase
            d.update()
            # detect transition INTO SPIN = impact moment
            if prev_phase == Die.PHASE_FLYING and d.phase == Die.PHASE_SPIN:
                self.shake_timer = 5
                for _ in range(8):
                    self.particles.append(Particle(
                        d.x, d.y,
                        random.choice([C_PURPLE, C_LAVENDER, (200, 160, 255)]),
                        size=random.uniform(2, 5),
                        vx=random.uniform(-4, 4),
                        vy=random.uniform(-5, -1),
                        life=random.randint(15, 28),
                    ))
            # detect LAND = final settle flash
            if prev_phase == Die.PHASE_BOUNCE and d.phase == Die.PHASE_LAND:
                for _ in range(5):
                    self.particles.append(Particle(
                        d.x, d.y,
                        C_GOLD,
                        size=random.uniform(2, 4),
                        vx=random.uniform(-3, 3),
                        vy=random.uniform(-3, 0),
                        life=random.randint(10, 20),
                    ))
                # re-eval hand once every die has landed
                self._eval()
        
        
        if self.state == GameState.LEVEL_UP:
            # Chờ người chơi chọn thuộc tính nâng cấp
            pass

    def handle_event(self, event):
        # Handle map manager events first (F1=editor, F3=minimap, F2=save)
        self.map_manager.handle_event(event)

        if event.type == pygame.USEREVENT + 1:
            self.on_enemy_dead()
        elif event.type == pygame.USEREVENT + 2:
            self.on_enemy_attack()
        elif event.type == pygame.USEREVENT + 3:
            if self.state != GameState.GAME_OVER:
                self.reset_dice()

        elif event.type == pygame.MOUSEBUTTONDOWN and event.button == 1:
            if self.state == GameState.ROLLING:
                for i, d in enumerate(self.dice):
                    if not d.empty and d.get_rect().collidepoint(event.pos):
                        self.toggle_lock(i)
                        break

        elif event.type == pygame.KEYDOWN:
            if self.state == GameState.MENU:
                if self.show_questions:
                    if event.key == pygame.K_ESCAPE:
                        self.show_questions = False
                else:
                    if event.key == pygame.K_UP:
                        self.menu_selected = max(0, self.menu_selected - 1)
                    elif event.key == pygame.K_DOWN:
                        self.menu_selected = min(3, self.menu_selected + 1)  # 4 options: 0-3
                    elif event.key in (pygame.K_RETURN, pygame.K_KP_ENTER):
                        self.handle_menu_selection()
            elif self.state == GameState.GAME_OVER:
                if event.key == pygame.K_r:
                    self.reset()
            elif self.state == GameState.ROLLING:
                # Xử lý câu hỏi boss nếu đang hiển thị
                if self.boss_active and self.boss_question and not self.boss_question_answered:
                    if event.key == pygame.K_UP:
                        self.answer_selected = max(0, self.answer_selected - 1)
                    elif event.key == pygame.K_DOWN:
                        self.answer_selected = min(len(self.boss_question["options"]) - 1, self.answer_selected + 1)
                    elif event.key in (pygame.K_RETURN, pygame.K_KP_ENTER):
                        self.submit_boss_answer()
                elif event.key == pygame.K_SPACE:
                    self.handle_reroll()
                elif event.key in (pygame.K_RETURN, pygame.K_KP_ENTER):
                    self.handle_attack()
            elif self.state == GameState.LEVEL_UP:
                if event.key == pygame.K_UP:
                    self.selected_upgrade = max(0, self.selected_upgrade - 1)
                elif event.key == pygame.K_DOWN:
                    self.selected_upgrade = min(len(self.level_up_options) - 1, self.selected_upgrade + 1)
                elif event.key in (pygame.K_RETURN, pygame.K_KP_ENTER):
                    self.apply_level_up_option(self.selected_upgrade)

    def handle_menu_selection(self):
        if self.menu_selected == 0:  # Chơi mới
            self.state = GameState.ROLLING
            self.reset()
            self.show_questions = False
        elif self.menu_selected == 1:  # Chơi tiếp
            if self.has_save:
                # TODO: load save
                self.state = GameState.ROLLING
                self.show_questions = False
            else:
                # Same as new game for now
                self.state = GameState.ROLLING
                self.reset()
                self.show_questions = False
        elif self.menu_selected == 2:  # Tạo câu hỏi từ file
            file_path = open_file_chooser()
            if file_path:
                imported = import_questions_from_file(file_path)
                if imported:
                    self.all_questions = imported
                    if save_questions(imported):
                        self.import_status = f"Thành công! Import {len(imported)} câu hỏi"
                    else:
                        self.import_status = "Lỗi lưu file"
                else:
                    self.import_status = "Không thể đọc file hoặc file không hợp lệ"
                self.import_status_timer = 180  # 3 giây
        elif self.menu_selected == 3:  # Thoát
            pygame.event.post(pygame.event.Event(pygame.QUIT))

    def draw(self):
        # screen shake
        shake_x, shake_y = 0, 0
        if self.shake_timer > 0:
            self.shake_timer -= 1
            shake_x = random.randint(-3, 3)
            shake_y = random.randint(-2, 2)

        draw_surf = screen
        if shake_x or shake_y:
            # blit everything onto a temp surface then offset
            draw_surf = pygame.Surface((SCREEN_W, SCREEN_H))

        if self.state == GameState.MENU:
                # Vẽ tilemap TRƯỚC enemy/player
                self.map_manager.draw(
                    draw_surf,
                    player_world_pos=(self.player.x, self.player.y) if self.player else None
                )
                # Giữ lại draw_background() nếu muốn kết hợp (castle, bats, cây...)
                draw_menu(draw_surf, self.menu_selected, self.has_save, self)
        else:
            # Draw game elements
            draw_background(draw_surf, self.frame, self.path_offset)
            self.map_manager.draw(
                draw_surf,
                player_world_pos=(self.player.x, self.player.y) if self.player else None
            )
            self.enemy.draw(draw_surf, self.player)
            self.player.draw(draw_surf)

            # Vẽ projectiles
            for proj in self.enemy.projectiles:
                proj.draw(draw_surf)

            for p in self.particles:
                p.draw(draw_surf)
            for ft in self.float_texts:
                ft.draw(draw_surf)

            draw_hud(draw_surf, self.player, self.level_count)

            # Wave indicator
            map_name = self.map_themes[self.current_map_index].upper()
            wave_label = f"BOSS FIGHT" if getattr(self.enemy, 'boss', False) else f"Wave {self.level_count}"
            wave_txt = font_sm.render(f"{map_name} — {wave_label}", True, C_LAVENDER)
            draw_surf.blit(wave_txt, (SCREEN_W//2 - wave_txt.get_width()//2, 55))

            draw_dice_panel(draw_surf, self.dice, self.hand_name, self.damage, self.rerolls, self.max_rerolls, self.combo_count, self.combo_effect)

            # Hiển thị câu hỏi khi fight boss
            if self.boss_active and self.boss_question and not self.boss_question_answered:
                draw_boss_question(draw_surf, self, self.answer_selected)

            if self.state == GameState.LEVEL_UP:
                draw_level_up_menu(draw_surf, self.player.level, self.level_up_options, self.selected_upgrade)
            if self.state == GameState.GAME_OVER:
                draw_game_over(draw_surf)

        if draw_surf is not screen:
            screen.fill((0, 0, 0))
            screen.blit(draw_surf, (shake_x, shake_y))


# ==================== ENTRY POINT ====================

def main():
    game = Game()

    print("=" * 50)
    print("  DICE ROGUELITE - Pygame")
    print("=" * 50)
    print("  SPACE  : Tung lại xúc xắc")
    print("  ENTER  : Tấn công kẻ thù")
    print("  Click  : Khóa/mở xúc xắc")
    print("  R      : Chơi lại (khi thua)")
    print("=" * 50)

    running = True
    while running:
        for event in pygame.event.get():
            if event.type == pygame.QUIT:
                running = False
            game.handle_event(event)

        game.update()
        game.draw()
        pygame.display.flip()
        clock.tick(FPS)

    pygame.quit()
    sys.exit()


if __name__ == "__main__":
    main() 